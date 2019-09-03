import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
Document()
def data2Doc(_path,output):
    doc=Document()
    df = pd.read_csv(open(_path,encoding='utf-8'), lineterminator='\n', header=0)
    for i in range(1,len(df['content_html'])):
        title=df.get('title').iloc[i]
        p_total = doc.add_heading(level=1)
        r_total = p_total.add_run(title)
        r_total.font.bold = True
        r_total.font.size = Pt(18)
        r_total.font.name = u'微软雅黑'
        r_total._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        if type(df.get('content_html').iloc[i])!=str:
            continue
        soup=BeautifulSoup(df.get('content_html').iloc[i],'lxml')
        # context=str(df.get('content_html').iloc[i])[50:-6]
        for p in soup.select('p'):
            doc.add_paragraph(p.text)
        if i%100==0:
            print("正在写入%s篇文章"%(str(i)))
    doc.save(output)
        # str_content='## %s\n %s\n'%(title,context)
        # doc.add_paragraph("正文")
        # with open(output,'a',encoding='utf-8') as f:
        #     f.write(str_content)

if __name__ == '__main__':
    data2Doc('./caoz的梦呓/article.list.csv','./caoz的梦呓/article.list.docx')
    # pd.DataFrame(['asdf','vvvv']).to_csv('aa.csv')
    # pd.DataFrame(pd.read_csv('./caoz的梦呓//aa.csv', encoding='gbk',lineterminator='\n', header=0))
